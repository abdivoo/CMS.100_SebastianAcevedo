#!/usr/bin/env python3
"""
parse_nexis_docx.py

Reads a Nexis‐exported .docx file containing a year's worth of articles,
correctly splits on article boundaries, and extracts for each article:
  - title
  - body text

Outputs a CSV with columns: year, title, body.
"""

import re
import pandas as pd
from docx import Document
import os

# ── CONFIG ──────────────────────────────────────────────────────────────────────

# Path to your .docx file
DOCX_PATH = "refined_methodology/ethical_ai_2025.DOCX"

# Year (can also parse from filename if you prefer)
YEAR = 2025

# Improved patterns
TITLE_PATTERN = re.compile(r"^\d+\s+of\s+\d+\s+DOCUMENTS?\s*$", re.IGNORECASE)
END_MARKER = "End of Document"
LENGTH_PATTERN = re.compile(r"^LENGTH:\s*(\d+)\s+words?", re.IGNORECASE)

# ── FUNCTIONS ───────────────────────────────────────────────────────────────────

def load_docx(path):
    """Load and return full text from a .docx as a single string with proper line breaks."""
    try:
        doc = Document(path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        return "\n".join(full_text)
    except Exception as e:
        print(f"Error loading document: {e}")
        return ""

def split_articles(raw_text):
    """Split the raw text into individual articles using more reliable markers."""
    # Split by End of Document marker
    article_chunks = []
    chunks = raw_text.split(END_MARKER)
    
    for chunk in chunks:
        chunk = chunk.strip()
        if not chunk:
            continue
            
        # Further split if there are multiple articles in one chunk
        lines = chunk.splitlines()
        current_article = []
        
        for i, line in enumerate(lines):
            current_article.append(line)
            
            # Check if we've found a new article start marker
            if i > 0 and TITLE_PATTERN.match(line) and len(current_article) > 5:
                # Found a new article, save the current one
                article_text = "\n".join(current_article[:-1])
                if article_text.strip():
                    article_chunks.append(article_text)
                # Start a new article
                current_article = [line]
        
        # Add the last article in the chunk
        if current_article:
            article_text = "\n".join(current_article)
            if article_text.strip():
                article_chunks.append(article_text)
    
    return article_chunks

def extract_article_info(article_text):
    """Extract title and body from an article chunk."""
    lines = article_text.splitlines()
    
    # Initialize variables
    title = None
    body_lines = []
    in_body = False
    capture_title = False
    
    for i, line in enumerate(lines):
        # Find the title (usually appears after the document count marker)
        if TITLE_PATTERN.match(line):
            # Title should be in the next non-empty line
            for j in range(i+1, min(i+10, len(lines))):
                if lines[j].strip() and not lines[j].startswith("LENGTH:"):
                    title = lines[j].strip()
                    break
            continue
            
        # Check for LENGTH marker which usually appears before the body
        if LENGTH_PATTERN.match(line):
            # Start capturing body text soon
            in_body = True
            continue
            
        # Capture body text (after LENGTH marker, excluding metadata)
        if in_body:
            # Skip headers and metadata sections
            if line.strip() in ["BODY", "BYLINE", "DATELINE", "HIGHLIGHT", "LOAD-DATE"]:
                continue
                
            # Skip lines starting with common metadata patterns
            if re.match(r"^(BYLINE|SECTION|LOAD-DATE|LANGUAGE|PUBLICATION-TYPE|JOURNAL-CODE):", line):
                continue
                
            # Add to body text if not empty and not metadata
            if line.strip() and not line.strip().startswith("Copyright"):
                body_lines.append(line)
    
    # Join body lines into a single string
    body = "\n".join(body_lines).strip()
    
    # Clean up the title if found
    if title:
        # Remove any publication names in brackets
        title = re.sub(r'\[.*?\]', '', title).strip()
    else:
        title = "Unknown Title"
        
    return title, body

# ── MAIN ────────────────────────────────────────────────────────────────────────

def main():
    if not os.path.exists(DOCX_PATH):
        print(f"Error: File {DOCX_PATH} not found. Please check the path.")
        return
        
    print(f"Loading DOCX file: {DOCX_PATH} …")
    raw = load_docx(DOCX_PATH)
    
    if not raw:
        print("Error: Failed to load or empty document.")
        return

    print("Splitting into article chunks…")
    chunks = split_articles(raw)
    print(f"Found {len(chunks)} article chunks.")
    
    if len(chunks) == 0:
        print("No articles found. Check if the document format matches the expected Nexis export format.")
        return

    records = []
    print("Extracting article information...")
    
    for i, chunk in enumerate(chunks):
        title, body = extract_article_info(chunk)
        
        # Skip if body is too short (likely not a real article)
        if len(body.split()) < 30:
            continue
            
        records.append({
            "year": YEAR,
            "title": title,
            "body": body
        })
        
        # Print progress
        if (i+1) % 10 == 0:
            print(f"Processed {i+1} articles...")

    df = pd.DataFrame(records)
    out_csv = f"nexis_{YEAR}_structured.csv"
    df.to_csv(out_csv, index=False)
    print(f"Done—exported {len(df)} articles to {out_csv}")
    
    # Print first few titles to verify
    if len(df) > 0:
        print("\nFirst 5 article titles for verification:")
        for i, title in enumerate(df["title"].head(5)):
            print(f"{i+1}. {title}")

if __name__ == "__main__":
    main()

# """
# parse_nexis_docx.py

# Reads a Nexis‐exported .docx file containing a year's worth of articles,
# correctly splits on article boundaries, and extracts for each article:
#   - title
#   - body text

# Outputs a CSV with columns: year, title, body.
# """

# import re
# import pandas as pd
# from docx import Document
# import os

# # ── CONFIG ──────────────────────────────────────────────────────────────────────

# # Path to your .docx file
# DOCX_PATH = "refined_methodology/ethical_ai_2021.DOCX"

# # Year (can also parse from filename if you prefer)
# YEAR = 2021

# # Improved patterns
# TITLE_PATTERN = re.compile(r"^\d+\s+of\s+\d+\s+DOCUMENTS?\s*$", re.IGNORECASE)
# END_MARKER = "End of Document"
# LENGTH_PATTERN = re.compile(r"^LENGTH:\s*(\d+)\s+words?", re.IGNORECASE)
# BODY_MARKER = re.compile(r"^BODY:?\s*$", re.IGNORECASE)
# PUBLICATION_PATTERN = re.compile(r"^[A-Z\s]+$")  # Often publication names are all caps

# # ── FUNCTIONS ───────────────────────────────────────────────────────────────────

# def load_docx(path):
#     """Load and return full text from a .docx as a single string with proper line breaks."""
#     try:
#         doc = Document(path)
#         full_text = []
#         for para in doc.paragraphs:
#             full_text.append(para.text)
#         return "\n".join(full_text)
#     except Exception as e:
#         print(f"Error loading document: {e}")
#         return ""

# def split_articles(raw_text):
#     """Split the raw text into individual articles."""
#     articles = []
#     current_article = []
#     in_article = False
    
#     lines = raw_text.splitlines()
    
#     for line in lines:
#         # New article starts with the document count marker
#         if TITLE_PATTERN.match(line):
#             if in_article and current_article:
#                 articles.append("\n".join(current_article))
#                 current_article = []
#             in_article = True
        
#         # End of article marker
#         if line.strip() == END_MARKER:
#             if in_article and current_article:
#                 articles.append("\n".join(current_article))
#                 current_article = []
#             in_article = False
#             continue
        
#         if in_article:
#             current_article.append(line)
    
#     # Add the last article if there is one
#     if in_article and current_article:
#         articles.append("\n".join(current_article))
    
#     return articles

# def extract_article_info(article_text):
#     """Extract title and body from an article chunk with improved title detection."""
#     lines = article_text.splitlines()
    
#     # Initialize variables
#     title = None
#     body_lines = []
#     in_body = False
#     title_candidates = []
    
#     # Step 1: Find potential title lines
#     for i, line in enumerate(lines):
#         if i < 15 and line.strip() and not line.startswith("LENGTH:") and not TITLE_PATTERN.match(line):
#             # Skip document count marker
#             if TITLE_PATTERN.match(line):
#                 continue
                
#             # Skip common metadata headers
#             if re.match(r"^(BYLINE|SECTION|LOAD-DATE|LANGUAGE|PUBLICATION-TYPE|JOURNAL-CODE):", line):
#                 continue
                
#             # Skip date lines (often start with a month name)
#             if re.match(r"^(January|February|March|April|May|June|July|August|September|October|November|December)\s+\d+,\s+\d{4}", line):
#                 continue
            
#             # Skip publication names (often all caps)
#             if PUBLICATION_PATTERN.match(line) and len(line.split()) <= 5:
#                 continue
                
#             # Add as title candidate
#             title_candidates.append(line.strip())
    
#     # Choose the best title candidate - usually the longest non-metadata line
#     title_candidates = [t for t in title_candidates if len(t) > 15 and t != "BODY"]
    
#     if title_candidates:
#         # Sort by length and pick the longest that's not too long
#         title_candidates.sort(key=len, reverse=True)
#         for candidate in title_candidates:
#             if len(candidate) <= 200:  # Reasonable title length
#                 title = candidate
#                 break
#         if not title:
#             title = title_candidates[0]  # Take first/longest if all are long
    
#     # Find body content
#     for i, line in enumerate(lines):
#         # Look for the BODY marker or LENGTH marker which often precedes body text
#         if BODY_MARKER.match(line) or LENGTH_PATTERN.match(line):
#             in_body = True
#             continue
        
#         if in_body:
#             # Skip metadata headers
#             if re.match(r"^(BYLINE|SECTION|LOAD-DATE|LANGUAGE|GRAPHIC|PUBLICATION-TYPE|JOURNAL-CODE):", line):
#                 continue
                
#             # End collecting at copyright notice
#             if line.strip().startswith("Copyright "):
#                 break
                
#             body_lines.append(line)
    
#     # Join body lines into a single string
#     body = "\n".join(body_lines).strip()
    
#     # If no title was found, mark as unknown
#     if not title:
#         title = "Unknown Title"
        
#     return title, body

# # ── MAIN ────────────────────────────────────────────────────────────────────────

# def main():
#     if not os.path.exists(DOCX_PATH):
#         print(f"Error: File {DOCX_PATH} not found. Please check the path.")
#         return
        
#     print(f"Loading DOCX file: {DOCX_PATH}...")
#     raw = load_docx(DOCX_PATH)
    
#     if not raw:
#         print("Error: Failed to load or empty document.")
#         return

#     print("Splitting into article chunks...")
#     article_chunks = split_articles(raw)
#     print(f"Found {len(article_chunks)} article chunks.")
    
#     if len(article_chunks) == 0:
#         print("No articles found. Check if the document format matches the expected Nexis export format.")
#         return

#     records = []
#     print("Extracting article information...")
    
#     for i, chunk in enumerate(article_chunks):
#         title, body = extract_article_info(chunk)
        
#         # Skip if body is too short (likely not a real article)
#         if len(body.split()) < 30:
#             continue
            
#         records.append({
#             "year": YEAR,
#             "title": title, 
#             "body": body
#         })
        
#         # Print progress
#         if (i+1) % 10 == 0 or i == 0:
#             print(f"Processed {i+1} articles...")
#             if i == 0:
#                 print(f"Sample title: {title[:80]}...")

#     df = pd.DataFrame(records)
#     out_csv = f"nexis_{YEAR}_structured.csv"
#     df.to_csv(out_csv, index=False)
#     print(f"Done—exported {len(df)} articles to {out_csv}")
    
#     # Print first few titles to verify
#     if len(df) > 0:
#         print("\nFirst 5 article titles for verification:")
#         for i, title in enumerate(df["title"].head(5)):
#             print(f"{i+1}. {title}")

# if __name__ == "__main__":
#     main()